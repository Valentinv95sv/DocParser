from docx.shared import Inches
from mailmerge import MailMerge
import datetime
import docx

class docx_insert_data_class:

    #метод, генерирующий название нового сгенерированного файла
    def get_file_name(self):
        filename = 'Blank_' + str(datetime.datetime.now().strftime("%H%M%S")) + '.docx'
        return filename

    #метод, возвращиющий все специальные поля в документе
    def get_fields(self, doc_path):
        template = doc_path
        document = MailMerge(template)
        return document.get_merge_fields()

    #метод по добавлению данных только в поля
    def merge_file_fields(self, doc_path, inserted_fields, filename):
        dict = inserted_fields
        document = MailMerge(doc_path)
        document.merge_pages([dict])
        document.write(filename)
        return filename

    #метод по добавлению данных только в таблицы
    def merge_file_tables(self, doc_path, inserted_tables, filename):
        document = MailMerge(doc_path)
        for i in inserted_tables:
            for j in i:
                a = list(j.keys())[1]
            document.merge_rows(a, i)
        document.write(filename)
        return filename

    #метод по добавлению данных в поля и поля таблиц, размеченных в документе
    def merge_file_fields_tables(self, doc_path, inserted_fields, inserted_tables2, filename):
        document = MailMerge(doc_path)
        for i in inserted_tables2:
            for j in i:
                a = list(j.keys())[0]
            document.merge_rows(a, i)
        document.merge_pages([inserted_fields])
        document.write(filename)
        return filename

#метод по добавлению изображений в формате png в поля таблиц, содержащих определенные символы
    def add_picture(self, docpath, picPath):
        doc = docx.Document(docpath)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key, value in picPath.items():
                            if ( paragraph.text == key):
                                paragraph.text = paragraph.text.replace(key, "")
                                a = paragraph.add_run()
                                a.add_picture(value, width=Inches(0.7), height=Inches(0.5))
        doc.save(docpath)




